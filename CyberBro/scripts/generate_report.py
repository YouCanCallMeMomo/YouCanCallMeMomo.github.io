import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Initialize document
doc = Document()

# Add Title
title = doc.add_heading('Laporan Akhir: SiagaSiber Youth System', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add Introduction
doc.add_heading('1. Pendahuluan', level=1)
p1 = doc.add_paragraph(
    "SiagaSiber Youth System adalah sebuah platform edukasi keamanan siber (Cyber Security) interaktif "
    "yang dirancang dengan pendekatan gamifikasi. Platform ini bertujuan untuk mengedukasi siswa/siswi SMP "
    "dan SMA di Indonesia agar memahami ancaman siber seperti Phishing, Web Defacement, Password Cracking, "
    "Forensik Digital, dan SQL Injection dengan cara yang aman dan menyenangkan."
)

doc.add_heading('2. Spesifikasi Teknologi', level=1)
doc.add_paragraph("Platform ini dibangun menggunakan teknologi mutakhir:", style='List Bullet')
doc.add_paragraph("Frontend Framework: React 18 & TypeScript", style='List Bullet')
doc.add_paragraph("Build Tool: Vite", style='List Bullet')
doc.add_paragraph("Styling: TailwindCSS", style='List Bullet')
doc.add_paragraph("Icons & Animations: Lucide React & Framer Motion", style='List Bullet')

doc.add_heading('3. Dokumentasi Fitur Aplikasi', level=1)

screenshot_dir = os.path.join(os.path.dirname(__file__), '../public/screenshots')

features = [
    {
        'file': '01_LoginPage.png',
        'title': 'Halaman Registrasi / Login',
        'desc': 'Cukup masukkan nama lengkap, alokasi password, dan email untuk masuk ke dalam sistem. Data disimpan persisten ke memori lokal.'
    },
    {
        'file': '02_Dashboard.png',
        'title': 'Dashboard / Landing Page',
        'desc': 'Beranda utama yang menampilkan ringkasan skor poin, daftar badge, dan akses navigasi cepat ke seluruh menu pembelajaran dan mini-games.'
    },
    {
        'file': '03_Modules.png',
        'title': 'Modul Pembelajaran',
        'desc': 'Materi singkat dan padat mengenai Phishing, Malware, Password, hingga Wi-Fi Publik.'
    },
    {
        'file': '04_LeakedVault.png',
        'title': 'Mini Game: The Leaked Vault',
        'desc': 'Mode permainan di mana pengguna diajak memilah database password yang bocor untuk membedakan password Kuat dan password Lemah.'
    },
    {
        'file': '05_Defacement.png',
        'title': 'Mini Game: Defacement Scanner',
        'desc': 'Permainan memulihkan tampilan situs sekolah yang diretas (Defacement) dengan mencari elemen-elemen palsu yang ditanam sang peretas.'
    },
    {
        'file': '06_ChatTrap.png',
        'title': 'Mini Game: The Chat Trap',
        'desc': 'Permainan menganalisis pesan percakapan untuk menentukan apakah itu aman atau merupakan pesan tipuan Phishing.'
    },
    {
        'file': '07_Forensic.png',
        'title': 'Mini Game: Forensic Investigator',
        'desc': 'Implementasi baru dalam sistem di mana pemain menelaah potongan data bukti digital (Forensik) untuk menentukan apakah log tersebut Aman atau Malicious (Anomali).'
    },
    {
        'file': '08_SqliScanner.png',
        'title': 'Mini Game: SQLi Scanner',
        'desc': 'Implementasi baru yang bertindak sebagai firewall game. Pengguna harus mengidentifikasi input normal vs. input injeksi SQL dengan cepat.'
    },
    {
        'file': '09_Quiz.png',
        'title': 'Kuis Dinamis Terstruktur',
        'desc': 'Sistem kuis dengan 600 bank soal yang menyediakan pilihan Mode Kuis: General, SQL Injection, atau Forensic.'
    },
    {
        'file': '10_Leaderboard.png',
        'title': 'Sistem Leaderboard & Badges',
        'desc': 'Mendata seluruh nilai pengguna lengkap dengan koleksi lencana/badge khusus jika pengguna menyelesaikan game atau mendapatkan skor sempurna di Kuis.'
    },
    {
        'file': '11_CyberBro.png',
        'title': 'Asisten AI CyberBro',
        'desc': 'Chatbot interaktif bersenjatakan integrasi API OpenAI (GPT-3.5/GPT-4o) untuk berkonsultasi seputar dunia keamanan siber.'
    }
]

for feature in features:
    img_path = os.path.join(screenshot_dir, feature['file'])
    doc.add_heading(feature['title'], level=2)
    doc.add_paragraph(feature['desc'])
    
    if os.path.exists(img_path):
        try:
            doc.add_picture(img_path, width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[Gambar gagal dimuat: {e}]", style='Intense Quote')
    else:
        doc.add_paragraph("[Gambar screenshot belum diambil / tidak ditemukan]", style='Intense Quote')

doc.add_heading('4. Kesimpulan', level=1)
doc.add_paragraph(
    "Dengan penerapan mini-games yang menguji ketelitian dalam klasifikasi ancaman siber "
    "(mulai dari Phishing, Defacement, hingga SQLi dan Forensik), aplikasi SiagaSiber dapat "
    "menjadi alat peraga canggih namun ringan bagi pendidik keamanan data. Semuanya tidak "
    "hanya diajarkan lewat teori, namun dinilai secara konkret melewati poin dan peraihan badge."
)

output_path = os.path.join(os.path.dirname(__file__), '../LAPORAN_SiagaSiber_Final.docx')
doc.save(output_path)
print(f"Laporan berhasil dibuat di: {output_path}")
