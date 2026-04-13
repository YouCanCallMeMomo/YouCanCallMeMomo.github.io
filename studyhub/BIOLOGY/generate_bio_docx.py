from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document()

# Default styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Margin setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_heading(text, level, color=RGBColor(0x2C, 0x3E, 0x50)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    return h

def add_paragraph(text, bold=False, italic=False, align=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    return p

def add_list_item(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.25 + level * 0.6)
    
    # Handle bold parts based on simple splitting if contains ':'
    if ':' in text:
        parts = text.split(':', 1)
        run_bold = p.add_run(parts[0] + ':')
        run_bold.bold = True
        run_bold.font.name = 'Calibri'
        run_norm = p.add_run(parts[1])
        run_norm.font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.name = 'Calibri'
    
    return p

# Cover / Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("Materi UTS Computational Biology")
title_run.font.size = Pt(24)
title_run.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
title_run.bold = True

add_paragraph("Ujian Tengah Semester - Ringkasan Materi & Persiapan", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# Metadata Section
doc.add_paragraph()
meta_p = doc.add_paragraph()
meta_p.add_run("Ketentuan Ujian\n").bold = True
meta_p.add_run("- Sifat: Close Book, Close Notes\n")
meta_p.add_run("- Alat Bantu: No Calculator")
meta_p.paragraph_format.left_indent = Cm(0.5)

doc.add_paragraph()
hr = doc.add_paragraph()
hr.add_run("_" * 70).font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7)

# SECTION 1
add_heading("1. Central Dogma of Molecular Biology & Mutations", 1)

add_heading("A. Coding & Template Strands", 2)
add_paragraph("DNA memiliki dua untai (strands) yang anti-paralel dan saling komplementer.")
add_list_item("Template Strand (Non-coding / Antisense strand): Untai DNA yang bertindak sebagai cetakan bagi enzim RNA polimerase selama proses transkripsi. Untai ini dibaca dari arah 3' ke 5', sehingga menghasilkan mRNA dengan arah 5' ke 3'.")
add_list_item("Coding Strand (Sense strand): Untai DNA pasangan dari template strand. Untai ini tidak digunakan sebagai cetakan, namun arah dan urutan nukleotidanya sama persis dengan mRNA yang dihasilkan (kecuali T diganti U).")

add_heading("B. Transcription & Translation", 2)
add_list_item("Transcription (Transkripsi): Proses mengubah informasi genetik dari DNA menjadi mRNA. Berlokasi di Nukleus (eukariota). Dikerjakan oleh RNA Polimerase.")
add_list_item("Translation (Translasi): Proses penerjemahan mRNA menjadi rangkaian asam amino pembentuk protein. Berlokasi di Ribosom. Melibatkan Kodon (triplet nukleotida di mRNA) dan tRNA pembawa asam amino.")

add_heading("C. Mutation Types & Effects", 2)
add_paragraph("Mutasi adalah perubahan pada urutan nukleotida DNA. Jenisnya:")
p_mut1 = doc.add_paragraph(style='List Bullet')
r1 = p_mut1.add_run("Point Mutation (Substitusi):")
r1.bold = True
add_list_item("Silent Mutation: Tidak mengubah urutan asam amino. Efek tidak ada.", level=1)
add_list_item("Missense Mutation: Mengubah satu asam amino menjadi yang lain. Efeknya mengubah struktur protein (contoh: Anemia Sel Sabit).", level=1)
add_list_item("Nonsense Mutation: Mengubah kodon menjadi kodon stop. Efek sintesis terganggu prematur (protein non-fungsional).", level=1)
p_mut2 = doc.add_paragraph(style='List Bullet')
r2 = p_mut2.add_run("Frameshift Mutation (Insersi / Delesi):")
r2.bold = True
add_paragraph("   Perubahan jumlah basa yang bukan kelipatan 3. Menggeser kerangka baca (reading frame) ribosom, sangat merusak seluruh kodon asam amino setelah mutasi.", italic=True)

doc.add_page_break()

# SECTION 2
add_heading("2. Basic Gene Expression", 1)

add_heading("A. Gene Expression", 2)
add_paragraph("Proses di mana informasi dari gen secara nyata digunakan dalam sintesis produk gen yang fungsional (protein/RNA non-coding).")

add_heading("B. RNA Splicing (Exons & Introns)", 2)
add_paragraph("Pada eukariota, hasil transkripsi awal pre-mRNA mengandung sekuens yang perlu dimodifikasi:")
add_list_item("Exons (Expressed sequences): Bagian gen yang mengkode protein.")
add_list_item("Introns (Intervening sequences): Bagian gen yang tidak mengkode protein.")
add_list_item("RNA Splicing: Proses memotong intron dan menyambungkan ekson yang difasilitasi spliceosome. ")
add_list_item("Alternative Splicing: Fenomena di mana 1 pre-mRNA yang sama disambungkan pada ekson yang berbeda-beda untuk memproduksi lebih dari 1 macam protein.")

# SECTION 3
add_heading("3. Bioinformatics Foundations", 1)
add_paragraph("Aplikasi ilmu komputer untuk memanage dan menganalisa data besar biologi molekuler.")

add_heading("A. Biological Databases (NCBI, GenBank)", 2)
add_list_item("Biological Database: Arsip digital online besar (DNA, Protein, Publikasi).")
add_list_item("NCBI (National Center for Biotechnology Information): Institusi krusial penyedia layanan publik data dan software alat biogenetik Amerika.")
add_list_item("GenBank: Database urutan nukleotida primer secara global oleh NCBI. Entri sekuens memiliki identitas unik berupa Accession Number.")

add_heading("B. Aplikasi Biotechnology Dasar", 2)
add_list_item("Sequence Alignment (BLAST): Mengukur tingkat kemiripan sekuens sampel misterius kita dengan seluruh data di database untuk mengidentifikasi organisme/fungsi patogen.")
add_list_item("Desain Primer PCR: Merancang struktur pemancing DNA spesifik di bioinformatika sebelum diproduksi nyata di laboratorium bioteknologi.")
add_list_item("Filogenetik: Menganalisa riwayat evolusioner dari relasi struktur rantai DNA makhluk tersebut dengan patogen masa silam.")

# SECTION 4
add_heading("4. Studi Kasus Bioinformatics", 1)
bg = parse_xml(r'<w:shd {} w:fill="EAF2F8" w:val="clear"/>'.format(nsdecls('w')))
p_case = doc.add_paragraph()
p_case._p.get_or_add_pPr().append(bg)
p_case.add_run("SKENARIO: Investigasi Wabah Penyakit\n").bold = True
p_case.add_run("Terdapat sampel patogen dari sebuah wabah akut di peternakan. Anda berperan sebagai peneliti bioinformatika.\n\n")
p_case.add_run("Langkah Penyelesaian:\n").bold = True
case_steps = [
    "1. Pencarian Database (BLAST): Memasukkan isolat primer DNA DNA patogen ke alat BLAST NCBI yang bersanding ke GenBank. Jika didapati hasil Identity 96% ke keluarga Avian Influenza, kemungkinan ini adalah wabah flu unggas varian baru.",
    "2. Coding Strand Check: Translasi dari pembacaan DNA dan RNA untuk mengetahui jenis protein surface virus yang tercipta.",
    "3. Analisa Mutasi: Menyelidiki apakah virus lama mengalami Missense Mutation yang membuat struktur serang virusnya baru dan menembus antibodi ternak.",
    "4. Penelusuran Splicing: Kadang virus ini memakai sifat sel mamalia yang melakukan Alternative Splicing, untuk mengeksploitasi 1 sel mamalia melipatgandakan beberapa varian RNA virus di masa datang."
]
for step in case_steps:
    p_case.add_run(f"  {step}\n")


out_path = r"c:\SEMESTER4\BIOLOGY\UTS_COMPUTATIONAL_BIOLOGY_MATERI.docx"
doc.save(out_path)
print("Docx created!")
