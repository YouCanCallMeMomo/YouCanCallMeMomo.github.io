from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# Margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def heading(text, level, color=RGBColor(0x2C, 0x3E, 0x50)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    return h

def paragraph(text, bold=False, italic=False, color=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    return p

# TITLE
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Latihan Soal & Pembahasan Lengkap")
tr.font.size = Pt(22)
tr.font.color.rgb = RGBColor(0, 102, 204)
tr.bold = True
sub = doc.add_paragraph("Computational Biology - Persiapan UTS", style='Subtitle')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
hr = doc.add_paragraph()
hr.add_run("_" * 70).font.color.rgb = RGBColor(200, 200, 200)

doc.add_paragraph("Dokumen ini berisi materi detail soal dan jawaban (studi kasus) melengkapi silabus komputasi biologi.")

# SECTION TABEL KODON
heading("Panduan Translasi mRNA & Tabel Asam Amino", 1, color=RGBColor(0, 51, 102))

paragraph("Sebelum masuk ke soal, sangat penting untuk mampu menggunakan tabel kodon.", italic=True)
doc.add_paragraph("Cara membaca:", style='List Bullet')
doc.add_paragraph("Basa 1 (Kiri) -> Baris", style='List Bullet')
doc.add_paragraph("Basa 2 (Atas) -> Kolom", style='List Bullet')
doc.add_paragraph("Basa 3 (Kanan) -> Baris spesifik dalam sel", style='List Bullet')

p = doc.add_paragraph()
r = p.add_run("Kodon AUG adalah START codon (Methionine). Kodon UAA, UAG, UGA adalah STOP codon.")
r.bold = True
r.font.color.rgb = RGBColor(204, 0, 0)
p.paragraph_format.left_indent = Cm(1)

# IMAGE INSERTION
img_path = r"c:\SEMESTER4\BIOLOGY\codon_table.png"
if os.path.exists(img_path):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_img = p_img.add_run()
    r_img.add_picture(img_path, width=Inches(6.0))
    paragraph("Gambar 1. Tabel Kodon Asam Amino (Berdasarkan mRNA, arah 5' ke 3')", italic=True).alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    paragraph("[Gambar Tabel Kodon tidak ditemukan]", color=RGBColor(255, 0, 0))

# SECTION 1
heading("Topik 1: Central Dogma & Mutasi", 1, color=RGBColor(0, 51, 102))

heading("Studi Kasus 1", 2)
paragraph("Bila diberikan urutan DNA Coding Strand:", bold=True)
p_code = doc.add_paragraph("5' - ATG - CCA - GCA - TGA - 3'")
p_code.runs[0].font.name = 'Consolas'
p_code.paragraph_format.left_indent = Cm(1)

paragraph("Pertanyaan:", bold=True)
doc.add_paragraph("1. Tentukan urutan nukleotida Template Strand!", style='List Number')
doc.add_paragraph("2. Tentukan urutan nukleotida pada mRNA hasil transkripsi!", style='List Number')
doc.add_paragraph("3. Terjemahkan urutan mRNA ke dalam rantai asam amino menggunakan tabel!", style='List Number')

paragraph("Pembahasan:", bold=True, color=RGBColor(39, 174, 96))
p_a1 = doc.add_paragraph("1. Template strand anti-paralel dan komplementer.")
p_a1.add_run("\n   Coding:   5' - ATG - CCA - GCA - TGA - 3'").font.name = 'Consolas'
p_a1.add_run("\n   Template: 3' - TAC - GGT - CGT - ACT - 5'").font.name = 'Consolas'

p_a2 = doc.add_paragraph("2. mRNA disintesis dari template strand arah 5' ke 3'. Isinya identik dengan Coding Strand hanya T jadi U.")
p_a2.add_run("\n   mRNA:     5' - AUG - CCA - GCA - UGA - 3'").font.name = 'Consolas'

p_a3 = doc.add_paragraph("3. Translasi mRNA (pembacaan 3 nukleotida per kodon):")
p_a3.add_run("\n   AUG -> Met (Start)")
p_a3.add_run("\n   CCA -> Proline (Pro)")
p_a3.add_run("\n   GCA -> Alanine (Ala)")
p_a3.add_run("\n   UGA -> STOP")
p_a3.add_run("\n   Rantai Peptida: Met - Pro - Ala").bold = True

heading("Studi Kasus 2: Analisis Mutasi", 2)
p_mut = doc.add_paragraph()
p_mut.add_run("Merujuk pada rantai  ")
r_m = p_mut.add_run("5' - AUG - CCA - GCA - UGA - 3'")
r_m.font.name = 'Consolas'
p_mut.add_run("\nApa efek terhadap rantai protein bila terjadi mutasi pada basa ke-4 (C -> A)?")

paragraph("Pembahasan:", bold=True, color=RGBColor(39, 174, 96))
doc.add_paragraph("Kodon awal CCA (Proline). Basa pertama C bermutasi menjadi A.\nSehingga kodon ke-2 menjadi ACA. Kodon ACA membentuk asam amino Threonine (Thr).")
val = doc.add_paragraph("Kesimpulan: Rantai sekarang adalah Met - Thr - Ala - Stop. Perubahan satu asam amino disebut Missense Mutation.")
val.runs[0].bold = True

doc.add_page_break()

# SECTION 2
heading("Topik 2: Gene Expression & Splicing", 1, color=RGBColor(0, 51, 102))
paragraph("Soal Konsep:", bold=True)
paragraph("Gen A memiliki cetak baku setara >400 asam amino di DNA, tapi protein keluaran akhirnya hanya sepanjang 120 asam amino. Jelaskan!")
paragraph("Pembahasan:", bold=True, color=RGBColor(39, 174, 96))
p_ex = doc.add_paragraph("Karena adanya tahapan RNA Splicing pada eukariota. Gen dalam DNA mengandung segmen pengkode (Exons) dan pemisah kosong (Introns). RNA mentranskripsinya secara primitif menjadi pre-mRNA. Sebelum keluar ke sitoplasma, mesin pemotong Spliceosome mengeksisi dan membuang rantai tak berguna (Introns) lalu menggabungkan (Splicing) Exon yang ada, menghasilkan struktur mature mRNA yang lebih pendek dari versi DNA masternya.")

# SECTION 3
heading("Topik 3: Bioinformatics Case Study", 1, color=RGBColor(0, 51, 102))
paragraph("Sebuah studi menemukan lalat di pedalaman kebun. Ilmuwan mengisolasi DNA liurnya ke bentuk file digital komputasi berformat .fasta.")

doc.add_paragraph("1. Platform biokomputasi apa yang kita tuju untuk mendeteksi identitas organisme secara akurat?", style='List Number')
paragraph("Jawaban: NCBI (National Center for Biotechnology Information). Alat spesifik di dalamnya adalah BLAST (Basic Local Alignment Search Tool).", color=RGBColor(39, 174, 96)).paragraph_format.left_indent = Cm(1)

doc.add_paragraph("2. Parameter angka apa saja yang kita nilai saat layar menampilkan urutan ranking dari basis datanya?", style='List Number')
paragraph("Jawaban: Identity Percentage (seberapa persen kemiripan susunan patogen kita dan dari arsip lama) serta E-Value (angka yang mengukur peluang kebetulan, E-Value mendekati 0 sangat valid maknanya).", color=RGBColor(39, 174, 96)).paragraph_format.left_indent = Cm(1)

doc.add_paragraph("3. Jika identifikasi berhasil menemukan ini spesies langka Ceratitis capitata, bagaimana simulasi agar kita memahami evolusi leluhur serangga tersebut?", style='List Number')
paragraph("Jawaban: Memasukkan deretan asalnya bersama kerabat keluarga spesies lain ke instrumen Phylogenetic Analysis untuk memperoleh gambar skema silsilah percabangan evolusioner pohon Phylogenetics.", color=RGBColor(39, 174, 96)).paragraph_format.left_indent = Cm(1)

p_footer = doc.add_paragraph("\n\n-- Latihan UTS Computational Biology (Aplikasi Central Dogma & Bioinformatika) --\nUniversitas: BINUS University")
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.runs[0].font.size = Pt(9)
p_footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

doc.save(r"c:\SEMESTER4\BIOLOGY\LATIHAN_SOAL_COMP_BIO.docx")
print("DOCX created successfully.")
